"""Tests for backend.ingest.extract_text — text extraction for ingest
source files (PDF / DOCX / txt / md / org).

PDF fixtures are hand-built minimal PDFs (correct xref offsets and
indirect references in /Kids, which pdfium requires) so no extra
dependency is needed. The DOCX fixture is round-tripped through
python-docx, mirroring the desktop extractor's Markdown shape
(src-tauri/src/commands/fs.rs extract_docx_with_library).
"""

import pytest

from backend.core.file_service import FsError
from backend.ingest.extract_text import extract_text


def build_pdf(page_contents: list[bytes]) -> bytes:
    """Assemble a minimal single- or multi-page PDF with correct xref.

    Objects are laid out as: 1 Catalog, 2 Pages, then per page a Page
    object, a content stream, and a Type1 font (3 slots per page), so
    page N's object id is 3 + 3*(N-1).
    """
    parts = [b"%PDF-1.4\n"]
    offsets: list[int] = []

    def add_object(body: bytes) -> int:
        offsets.append(sum(len(p) for p in parts))
        parts.append(f"{len(offsets)} 0 obj\n".encode() + body + b"\nendobj\n")
        return len(offsets)

    add_object(b"<< /Type /Catalog /Pages 2 0 R >>")
    kids = " ".join(f"{3 + 3 * i} 0 R" for i in range(len(page_contents)))
    add_object(
        f"<< /Type /Pages /Kids [{kids}] /Count {len(page_contents)} >>".encode()
    )

    for content in page_contents:
        content_id = len(offsets) + 2
        font_id = content_id + 1
        add_object(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Contents {content_id} 0 R /Resources << /Font "
                f"<< /F1 {font_id} 0 R >> >> >>"
            ).encode()
        )
        stream = (
            f"<< /Length {len(content)} >>\nstream\n".encode() + content + b"\nendstream"
        )
        add_object(stream)
        add_object(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    xref_pos = sum(len(p) for p in parts)
    count = len(offsets) + 1
    parts.append(f"xref\n0 {count}\n".encode() + b"0000000000 65535 f \n")
    for off in offsets:
        parts.append(f"{off:010d} 00000 n \n".encode())
    parts.append(
        f"trailer\n<< /Size {count} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode()
    )
    return b"".join(parts)


TEXT_OP = b"BT /F1 24 Tf 72 720 Td (Hello World from PDF page one.) Tj ET"
SECOND_PAGE_OP = b"BT /F1 24 Tf 72 700 Td (Second page text here.) Tj ET"


class TestPdfExtraction:
    def test_single_page_pdf(self, tmp_path):
        path = tmp_path / "one.pdf"
        path.write_bytes(build_pdf([TEXT_OP]))
        assert extract_text(str(path)) == "Hello World from PDF page one."

    def test_multi_page_join_rule(self, tmp_path):
        # Documented join rule: pages joined with "\n\n", except "\n"
        # when the preceding page's text lacks a trailing newline.
        path = tmp_path / "two.pdf"
        path.write_bytes(build_pdf([TEXT_OP, SECOND_PAGE_OP]))
        out = extract_text(str(path))
        assert out == "Hello World from PDF page one.\nSecond page text here."

    def test_case_insensitive_extension(self, tmp_path):
        path = tmp_path / "two.PDF"
        path.write_bytes(build_pdf([TEXT_OP, SECOND_PAGE_OP]))
        assert "Second page text here." in extract_text(str(path))

    def test_corrupt_pdf_raises_fs_error(self, tmp_path):
        path = tmp_path / "corrupt.pdf"
        path.write_bytes(b"not a pdf at all")
        with pytest.raises(FsError, match="Failed to open PDF"):
            extract_text(str(path))


class TestDocxExtraction:
    def test_structured_markdown_round_trip(self, tmp_path):
        from docx import Document

        path = tmp_path / "doc.docx"
        doc = Document()
        doc.add_heading("Title Here", level=1)
        doc.add_paragraph("Some normal paragraph text.")
        p = doc.add_paragraph()
        r = p.add_run("bold text")
        r.bold = True
        r2 = p.add_run(" and italic")
        r2.italic = True
        doc.add_paragraph("")  # empty paragraph must be dropped
        doc.add_paragraph("bullet item", style="List Bullet")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B|escaped"
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"
        doc.save(str(path))

        out = extract_text(str(path))
        assert out == (
            "# Title Here\n\n"
            "Some normal paragraph text.\n\n"
            "**bold text*** and italic*\n\n"
            "- bullet item\n"
            "| A | B\\|escaped |\n"
            "| --- | --- |\n"
            "| C | D |\n\n"
        )

    def test_heading_levels_and_missing_style(self, tmp_path):
        from docx import Document

        path = tmp_path / "levels.docx"
        doc = Document()
        doc.add_heading("H1", level=1)
        doc.add_heading("H2", level=2)
        doc.save(str(path))
        assert extract_text(str(path)) == "# H1\n\n## H2\n\n"

    def test_unreadable_docx_raises_fs_error(self, tmp_path):
        path = tmp_path / "broken.docx"
        path.write_bytes(b"not a zip")
        with pytest.raises(FsError, match="Failed to parse DOCX"):
            extract_text(str(path))


class TestTextPassthrough:
    @pytest.mark.parametrize("ext", ["txt", "md", "markdown"])
    def test_text_files_passed_through(self, tmp_path, ext):
        path = tmp_path / f"note.{ext}"
        content = "# Note\n\nplain text body\n"
        path.write_text(content, encoding="utf-8")
        assert extract_text(str(path)) == content

    def test_missing_text_file_raises_fs_error(self, tmp_path):
        with pytest.raises(FsError, match="Failed to read file"):
            extract_text(str(tmp_path / "missing.txt"))


class TestOrgExtraction:
    def test_org_to_markdown(self, tmp_path):
        path = tmp_path / "notes.org"
        path.write_text(
            "#+TITLE: My Doc\n"
            "#+AUTHOR: Bob\n"
            "* Heading One\n"
            "** Sub Heading\n"
            "See [[target][desc]] and [[other]].\n"
            "#+BEGIN_SRC python\n"
            "print(1)\n"
            "#+END_SRC\n"
            "| a | b |\n"
            "|---+---|\n"
            "| 1 | 2 |\n",
            encoding="utf-8",
        )
        assert extract_text(str(path)) == (
            "# My Doc\n"
            "**Author:** Bob\n"
            "# Heading One\n"
            "## Sub Heading\n"
            "See [desc](target) and <other>.\n"
            "```python\n"
            "print(1)\n"
            "```\n"
            "| a | b |\n"
            "|---|---|\n"
            "| 1 | 2 |"
        )


class TestUnsupported:
    def test_unsupported_extension_raises_fs_error(self, tmp_path):
        path = tmp_path / "clip.rtf"
        path.write_text("x", encoding="utf-8")
        with pytest.raises(FsError, match="Unsupported source format: .rtf"):
            extract_text(str(path))
