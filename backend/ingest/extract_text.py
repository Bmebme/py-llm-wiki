"""Text extraction for ingest source files.

Port of the extraction pipelines the desktop app ran in
src-tauri/src/commands/fs.rs (preprocess_file at line ~134, plus the
extract_* helpers it calls) and src-tauri/src/commands/ebook.rs, scoped
to the formats the Python port supports today:

  .pdf       pypdfium2 per-page text extraction
  .docx      python-docx, mirroring the *spirit* of the desktop DOCX
             extractor (fs.rs extract_docx_with_library): structured
             Markdown with heading markers, list markers, bold/italic
             run formatting, and pipe tables
  .org       port of org_to_markdown (fs.rs ~253)
  .txt/.md/.markdown   plain text passthrough

All functions raise backend.core.file_service.FsError for user-facing
errors.
"""

from __future__ import annotations

import re
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from backend.core import file_service
from backend.core.file_service import FsError

# Text-like extensions passed straight through to the file contents.
TEXT_PASSTHROUGH_EXTENSIONS = frozenset({"txt", "md", "markdown"})

_SUPPORTED_EXTENSIONS = frozenset({"pdf", "docx", "org"}) | TEXT_PASSTHROUGH_EXTENSIONS

# The desktop's extract_pdf_text used the renderer's `page.text().all()`;
# pypdfium2's get_text_bounded() returns the same per-page text content.


def extract_text(path: str) -> str:
    """Extract plain/Markdown text from an ingestable source file.

    Returns the extracted text as a string. Raises FsError for missing
    files, unreadable files, unsupported extensions, and documents that
    cannot be parsed.
    """
    p = Path(path)
    ext = p.suffix.lstrip(".").lower()

    if ext in TEXT_PASSTHROUGH_EXTENSIONS:
        return file_service.read_text(p)

    if ext == "pdf":
        return _extract_pdf_text(p)

    if ext == "docx":
        return _extract_docx_text(p)

    if ext == "org":
        return _extract_org_text(p)

    raise FsError(
        f"Unsupported source format: .{ext} (supported: {', '.join(sorted(_SUPPORTED_EXTENSIONS))})"
    )


# ─── PDF ─────────────────────────────────────────────────────────────

def _extract_pdf_text(path: Path) -> str:
    """Extract per-page text via pypdfium2, skipping pages whose text
    extraction raises.

    Page joining rule (documented contract): pages are joined with
    ``\\n\\n``, except that the separator is ``\\n`` when the preceding
    page's text does not end with a newline — so the last word of one
    page never gets glued to the first word of the next. Page text is
    kept verbatim (including its own trailing newlines).
    """
    try:
        document = pdfium.PdfDocument(str(path))
    except Exception as exc:  # PDFium raises on missing/corrupt/encrypted files
        raise FsError(f"Failed to open PDF '{path}': {exc}") from exc

    chunks: list[str] = []
    try:
        for page in document:
            try:
                textpage = page.get_textpage()
            except Exception:
                continue  # per-page extraction error -> skip the page
            try:
                chunks.append(textpage.get_text_bounded())
            except Exception:
                continue
            finally:
                textpage.close()
    finally:
        document.close()

    out = ""
    for chunk in chunks:
        if not out:
            out = chunk
        elif out.endswith("\n"):
            out += "\n\n" + chunk
        else:
            out += "\n" + chunk
    return out


# ─── DOCX ────────────────────────────────────────────────────────────

def _extract_docx_text(path: Path) -> str:
    """Extract DOCX as structured Markdown via python-docx.

    Mirrors the desktop's extract_docx_with_library (fs.rs ~673) in
    spirit: heading styles become ``#``-prefixed headings (level from the
    style name's first digit), list paragraphs become ``-`` items, bold
    and italic runs are wrapped in ``**``/``*``/``***``, and tables become
    pipe tables with a ``| --- |`` header separator row. Normal
    paragraphs end with a blank line; headings end with one too. Cells
    escape ``|`` as ``\\|``. Empty paragraphs are dropped.
    """
    try:
        document = Document(str(path))
    except Exception as exc:
        raise FsError(f"Failed to parse DOCX '{path}': {exc}") from exc

    result: list[str] = []
    for block in document.iter_inner_content():
        if isinstance(block, Paragraph):
            para_text = ""
            for run in block.runs:
                text = run.text or ""
                if not text:
                    continue
                is_bold = bool(run.bold)
                is_italic = bool(run.italic)
                if is_bold and is_italic:
                    para_text += f"***{text}***"
                elif is_bold:
                    para_text += f"**{text}**"
                elif is_italic:
                    para_text += f"*{text}*"
                else:
                    para_text += text

            text = para_text.strip()
            if not text:
                continue

            style_name = block.style.name if block.style is not None else ""
            if "Heading" in style_name or "heading" in style_name:
                # Level from the style name's first digit, default 1
                # (fs.rs extract_docx_with_library: "Heading 2" -> 2).
                level = 1
                for ch in style_name:
                    if ch.isascii() and ch.isdigit():
                        level = int(ch)
                        break
                result.append(f"{'#' * level} {text}\n\n")
            elif _is_list_paragraph(block):
                result.append(f"- {text}\n")
            else:
                result.append(f"{text}\n\n")

        elif isinstance(block, Table):
            rows: list[list[str]] = []
            for row in block.rows:
                cells: list[str] = []
                for cell in row.cells:
                    cell_text = "".join(p.text for p in cell.paragraphs)
                    cells.append(cell_text.strip().replace("|", "\\|"))
                rows.append(cells)
            if rows:
                max_cols = max(len(r) for r in rows)
                for i, row in enumerate(rows):
                    padded = row + [""] * (max_cols - len(row))
                    result.append(f"| {' | '.join(padded)} |\n")
                    if i == 0:
                        result.append("|" + " --- |" * max_cols + "\n")
                result.append("\n")

    return "".join(result)


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    """List detection: direct numbering (numPr) or a "List …" style.

    The desktop checked the paragraph's numbering property only
    (fs.rs: para.property.numbering_property.is_some()); python-docx
    exposes that as a numPr on the paragraph XML. Style-based lists
    (e.g. "List Bullet") attach numbering at the style level, so the
    style name is checked as well — a deliberate, documented extension
    of the desktop check so `- item` markers still appear for them.
    """
    ppr = paragraph._p.pPr
    if ppr is not None and ppr.numPr is not None:
        return True
    if paragraph.style is not None:
        return "list" in paragraph.style.name.lower()
    return False


# ─── Org ─────────────────────────────────────────────────────────────

def _extract_org_text(path: Path) -> str:
    """Port of extract_org_text (fs.rs:240-251): read as text, then
    convert Org syntax to Markdown-shaped text (org_to_markdown)."""
    try:
        content = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise FsError("File is not valid UTF-8 text") from exc
    except OSError as exc:
        raise FsError(f"Failed to read Org file '{path}': {exc}") from exc
    return _org_to_markdown(content)


def _org_to_markdown(content: str) -> str:
    """Port of org_to_markdown (fs.rs:253-~330).

    #+BEGIN_SRC / #+BEGIN_EXAMPLE / #+BEGIN_QUOTE blocks become fenced
    code blocks (verbatim inside; a dangling open block is closed with
    ```), #+TITLE becomes `# title`, other keywords become `**Key:**
    value`, `* headings` become `#` headings, Org table separator lines
    are normalized, and [[target][description]] / [[target]] links become
    [description](target) / <target>.
    """
    normalized = content.replace("\r\n", "\n").replace("\r", "\n")
    # Rust `lines()` drops a single trailing empty segment; split("\n")
    # keeps it.
    raw_lines = normalized.split("\n")
    if raw_lines and raw_lines[-1] == "":
        raw_lines.pop()

    output: list[str] = []
    block_end: str | None = None

    for line in raw_lines:
        trimmed = line.strip()
        upper = trimmed.upper()

        if block_end is not None:
            if upper == block_end:
                output.append("```")
                block_end = None
            else:
                output.append(line)
            continue

        if upper.startswith("#+BEGIN_SRC"):
            # trimmed[11:] skips "#+BEGIN_SRC" + the separating space.
            language = trimmed[11:].strip().split()[0] if trimmed[11:].strip() else ""
            output.append(f"```{language}")
            block_end = "#+END_SRC"
            continue
        if upper == "#+BEGIN_EXAMPLE":
            output.append("```text")
            block_end = "#+END_EXAMPLE"
            continue
        if upper == "#+BEGIN_QUOTE":
            output.append("```text")
            block_end = "#+END_QUOTE"
            continue

        keyword = _parse_org_keyword(trimmed)
        if keyword is not None:
            key, value = keyword
            if key == "TITLE":
                output.append(f"# {value}")
            elif key not in ("OPTIONS", "PROPERTY", "SETUPFILE"):
                output.append(f"**{_title_case_ascii(key)}:** {value}")
            continue

        heading = _parse_org_heading(line)
        if heading is not None:
            level, heading_text = heading
            output.append(f"{'#' * min(level, 6)} {_convert_org_links(heading_text)}")
            continue

        if (
            trimmed.startswith("|")
            and trimmed.endswith("|")
            and "+" in trimmed
            and all(c in "|+-: " for c in trimmed)
        ):
            output.append(trimmed.replace("+", "|"))
            continue

        output.append(_convert_org_links(line))

    if block_end is not None:
        output.append("```")
    return "\n".join(output)


def _parse_org_keyword(line: str) -> tuple[str, str] | None:
    """Port of parse_org_keyword (fs.rs:~332-340)."""
    if not line.startswith("#+"):
        return None
    rest = line[2:]
    if ":" not in rest:
        return None
    key, value = rest.split(":", 1)
    if not key or not all(c.isascii() and (c.isalnum() or c == "_") for c in key):
        return None
    return (key.upper(), value.strip())


def _parse_org_heading(line: str) -> tuple[int, str] | None:
    """Port of parse_org_heading (fs.rs:~342-348): leading stars + space."""
    stars = 0
    while stars < len(line) and line[stars] == "*":
        stars += 1
    if stars == 0 or stars >= len(line) or line[stars] != " ":
        return None
    return (stars, line[stars + 1:].strip())


def _title_case_ascii(value: str) -> str:
    """Port of title_case_ascii (fs.rs:~350-357)."""
    lower = value.replace("_", " ").lower()
    if not lower:
        return ""
    return lower[0].upper() + lower[1:]


def _convert_org_links(line: str) -> str:
    """Port of convert_org_links (fs.rs:~359-376)."""
    out: list[str] = []
    rest = line
    while True:
        start = rest.find("[[")
        if start < 0:
            out.append(rest)
            break
        out.append(rest[:start])
        after = rest[start + 2:]
        end = after.find("]]")
        if end < 0:
            out.append(rest[start:])
            break
        inner = after[:end]
        if "][" in inner:
            target, description = inner.split("][", 1)
            out.append(f"[{description}]({target})")
        else:
            out.append(f"<{inner}>")
        rest = after[end + 2:]
    return "".join(out)
